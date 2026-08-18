#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将《应急通信装备测试报告》原始文档转换为固定模板。

用法：
    python3 scripts/make_test_report_template.py \
        --source "/path/to/应急通信装备测试报告-20260529(1).docx"
"""

from __future__ import annotations

import argparse
import copy
import pathlib
import sys

from docx import Document  # type: ignore
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.table import _Cell


PROJECT_ROOT = pathlib.Path(__file__).resolve().parent.parent
DEFAULT_OUTPUT = PROJECT_ROOT / "sources" / "测试报告" / "应急通信装备测试报告模板.docx"


# 段落下标 -> 新内容；每项为 (文本, 原段落运行下标)，用于保留原字体格式。
PARAGRAPH_TEMPLATE: dict[int, list[tuple[str, int]]] = {
    2: [
        (
            "为严格落实应急指挥通信保障工作部署，筑牢汛期安全通信防线，提升应急通信装备实战应用能力，"
            "保障极端天气及突发险情下指挥链路畅通、设备运行稳定、通信效能最大化发挥，"
            "近期【组织单位】组织开展应急通信装备专项测试工作。现将本次测试具体情况汇报如下：",
            0,
        )
    ],
    4: [
        ("测试范围：", 0),
        (
            "本次测试面向【覆盖区数量】个区开展，覆盖【覆盖单位层级】，"
            "重点对【设备类型清单】等【设备类型数量】类通信装备进行全面检测，共测试【测试设备总数】台设备。",
            1,
        ),
    ],
    5: [
        ("测试时间：", 0),
        ("【测试开始日期】—【测试结束日期】", 1),
    ],
    6: [
        ("测试要求：", 0),
        (
            "对现有常用通信装备进行全量测试，对库存备用装备进行摸底筛查、拉长周期测试，"
            "确保平时、战时有设备可用。",
            1,
        ),
    ],
    9: [
        (
            "卫星电话共【卫星电话总数】部：移动卫星电话【移动卫星电话总数】部"
            "（市局共【市局移动卫星电话数】部，大厅持有【大厅持有数】部、库存【库存数】部）"
            "和固移双用基座卫星电话【基座卫星电话总数】部（市局【市局基座卫星电话数】部）。",
            0,
        )
    ],
    10: [
        (
            "【不可测试说明】。本次实际测试【实际测试数】部，其中正常【正常数】部，异常【异常数】部。",
            0,
        )
    ],
    13: [
        (
            "目前【覆盖区数量】区370MHz手台总计【手台总数】部，剔除各区因自购无法与市局互联互通的设备后，"
            "共拉动测试【拉动测试数】部。已完成【覆盖区数量】区常用设备【常用设备测试数】部设备通信测试，"
            "其中正常【正常数】部、异常【异常数】部。",
            0,
        )
    ],
    17: [
        (
            "共拉动测试【覆盖区数量】个区布控球【布控球总数】台，"
            "其中正常接入【正常接入数】台，未接入【未接入数】台。",
            0,
        )
    ],
    20: [
        (
            "共拉动测试【覆盖区数量】个区单兵图传【单兵图传总数】台，"
            "其中正常接入【正常接入数】台，未接入【未接入数】台。"
            "暂无单兵图传的区：【暂无单兵图传区清单】。",
            0,
        )
    ],
    21: [("【单兵图传使用情况说明】", 0)],
    24: [
        (
            "共拉动测试【覆盖区数量】个区无人机【无人机总数】台，"
            "由于区局并未全都配备专业飞手，主要测试无人机的图传功能。"
            "其中正常接入【正常接入数】台，未接入【未接入数】台。"
            "暂无无人机的区：【暂无无人机区清单】。",
            0,
        )
    ],
    25: [("表5 无人机未接入清单", 0)],
    27: [
        (
            "共拉动测试【覆盖区数量】个区指挥通信车【指挥车总数】辆，测试过程中【整体测试情况】。"
            "其中能与市局建立通信的【正常数】辆，无法建立通信的【异常数】辆，"
            "暂无可联通指挥车的区：【暂无指挥车区清单】。",
            0,
        )
    ],
    28: [("表6 应急指挥通信车未建立通信清单", 0)],
    29: [("目前市局共有指挥通信车【市局指挥车数量】辆，明细清单见下：", 0)],
    34: [
        (
            "共拉动测试【覆盖区数量】个区会议平板，音视频效果【整体测试情况】。",
            0,
        )
    ],
    36: [
        (
            "共测试【覆盖区数量】个区应急叫应终端【叫应终端总数】台，正常【正常数】台，异常【异常数】台；"
            "暂无叫应终端的区：【无叫应终端区清单】。",
            0,
        )
    ],
    40: [
        ("1.存在问题：", 0),
        ("【问题描述与示例】", 2),
    ],
    41: [
        ("2.解决建议：", 0),
        ("【解决建议】", 2),
    ],
    43: [
        ("1.存在问题：", 0),
        ("【问题描述与示例】", 2),
    ],
    44: [
        ("2.解决建议：", 0),
        ("【解决建议】", 2),
    ],
    46: [
        ("1.存在问题：", 0),
        ("【问题描述与示例】", 2),
    ],
    47: [("2.解决建议：【建议内容】", 0)],
    48: [("（1）【建议一】", 0)],
    49: [("（2）【建议二】", 0)],
    52: [("【XXXX年XX月XX日】", 0)],
}


# 表格下标 -> 模板行方案。
# archetype_rows 为原始表行下标，生成时按原样式复制；rows 为占位内容。
TABLE_TEMPLATES: dict[int, dict] = {
    0: {
        "archetype_rows": [1, 2, 3, 4],
        "rows": [
            ["1", "【行政区（共X部，Y部故障）】", "【故障单位/地点】", "【数量】", "【故障原因】", "【维修进度】"],
            ["2", "", "【故障单位/地点】", "【数量】", "【故障原因】", "【维修进度】"],
            ["3", "【行政区（共X部，Y部故障）】", "【故障单位/地点】", "【数量】", "【故障原因】", "【维修进度】"],
            ["4", "", "【故障单位/地点】", "【数量】", "【故障原因】", "【维修进度】"],
        ],
        "total": ["合计", "合计", "合计", "【合计】", "-", "-"],
    },
    1: {
        "archetype_rows": [1, 1, 1],
        "rows": [
            ["1", "【行政区（共X部，Y部故障）】", "【故障单位/地点】", "【数量】", "【故障原因】", "【维修进度】"],
            ["2", "【行政区（共X部，Y部故障）】", "【故障单位/地点】", "【数量】", "【故障原因】", "【维修进度】"],
            ["3", "【行政区（共X部，Y部故障）】", "【故障单位/地点】", "【数量】", "【故障原因】", "【维修进度】"],
        ],
        "total": ["合计", "合计", "合计", "【合计】", "-", "-"],
    },
    2: {
        "archetype_rows": [1, 1, 1],
        "rows": [
            ["1", "【行政区】", "【持有数量】", "【未接入数量】", "【未接入原因】"],
            ["2", "【行政区】", "【持有数量】", "【未接入数量】", "【未接入原因】"],
            ["3", "【行政区】", "【持有数量】", "【未接入数量】", "【未接入原因】"],
        ],
        "total": ["合计", "合计", "-", "【合计】", "-"],
    },
    3: {
        "archetype_rows": [1, 1, 1],
        "rows": [
            ["1", "【行政区】", "【持有数量】", "【未接入数量】"],
            ["2", "【行政区】", "【持有数量】", "【未接入数量】"],
            ["3", "【行政区】", "【持有数量】", "【未接入数量】"],
        ],
        "total": ["合计", "合计", "-", "【合计】"],
    },
    4: {
        "archetype_rows": [1, 1, 1],
        "rows": [
            ["1", "【行政区】", "【持有数量】", "【未接入数量】", "【未接入原因】"],
            ["2", "【行政区】", "【持有数量】", "【未接入数量】", "【未接入原因】"],
            ["3", "【行政区】", "【持有数量】", "【未接入数量】", "【未接入原因】"],
        ],
        "total": ["合计", "合计", "-", "【合计】", "-"],
    },
    5: {
        "archetype_rows": [1, 1, 1],
        "rows": [
            ["1", "【行政区】", "【无法建立通信数量】", "【原因】"],
            ["2", "【行政区】", "【无法建立通信数量】", "【原因】"],
            ["3", "【行政区】", "【无法建立通信数量】", "【原因】"],
        ],
        "total": ["合计", "合计", "【合计】", "-"],
    },
    6: {
        "archetype_rows": [1, 2],
        "rows": [
            ["1", "【车辆名称】", "【品牌】", "【型号】", "【数量】", "【主要设备】"],
            ["2", "【车辆名称】", "【品牌】", "【型号】", "【数量】", "【主要设备】"],
        ],
        "total": None,
    },
    8: {
        "archetype_rows": [2, 3, 4, 5, 6, 7],
        "rows": [
            ["1", "【行政区（共X部，Y部故障）】", "【故障单位/地点】", "【数量】", "【故障原因】", "【维修进度】"],
            ["2", "", "【故障单位/地点】", "【数量】", "【故障原因】", "【维修进度】"],
            ["3", "【行政区（共X部，Y部故障）】", "【故障单位/地点】", "【数量】", "【故障原因】", "【维修进度】"],
            ["4", "", "【故障单位/地点】", "【数量】", "【故障原因】", "【维修进度】"],
            ["5", "【行政区（共X部，Y部故障）】", "【故障单位/地点】", "【数量】", "【故障原因】", "【维修进度】"],
            ["6", "", "【故障单位/地点】", "【数量】", "【故障原因】", "【维修进度】"],
        ],
        "total": ["合计", "合计", "合计", "【合计】", "-", "-"],
    },
}


def _run_properties(run) -> object | None:
    """返回运行属性元素的深拷贝，供新运行复用字体格式。"""
    rpr = run._r.find(qn("w:rPr"))
    return copy.deepcopy(rpr) if rpr is not None else None


def set_paragraph(paragraph, segments: list[tuple[str, int]]) -> None:
    """按段重写文本，保留每个段落在原文档中的字体格式。"""
    runs = list(paragraph.runs)
    rprs = [_run_properties(run) for run in runs]
    for run in runs:
        run._r.getparent().remove(run._r)
    for text, source_index in segments:
        run = paragraph.add_run(text)
        if source_index < len(rprs) and rprs[source_index] is not None:
            run._r.insert(0, copy.deepcopy(rprs[source_index]))
        if "【" in text:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def set_cell_text(cell, text: str) -> None:
    """重写单元格文本，保留原单元格首个运行的字体格式。"""
    paragraphs = cell.paragraphs
    for paragraph in paragraphs[1:]:
        paragraph._p.getparent().remove(paragraph._p)
    paragraph = paragraphs[0]
    runs = list(paragraph.runs)
    rprs = [_run_properties(run) for run in runs]
    for run in runs:
        run._r.getparent().remove(run._r)
    run = paragraph.add_run(text)
    if rprs and rprs[0] is not None:
        run._r.insert(0, copy.deepcopy(rprs[0]))
    if "【" in text:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def set_row_texts(table, row, texts: list[str]) -> None:
    """按列写入行文本；跳过纵向合并单元格的延续格。"""
    row_element = row._tr if hasattr(row, "_tr") else row
    tc_list = row_element.findall(qn("w:tc"))
    if len(tc_list) != len(texts):
        raise ValueError(f"行单元格数量 {len(tc_list)} 与占位内容数量 {len(texts)} 不一致")
    for tc, text in zip(tc_list, texts):
        tc_pr = tc.find(qn("w:tcPr"))
        v_merge = tc_pr.find(qn("w:vMerge")) if tc_pr is not None else None
        if v_merge is not None and v_merge.get(qn("w:val")) is None:
            continue
        set_cell_text(_Cell(tc, table), text)


def rebuild_table(table, spec: dict) -> None:
    """按原表头、合计行与指定样例行重建模板表格。"""
    tbl = table._tbl
    tr_list = list(tbl.findall(qn("w:tr")))
    header = tr_list[0]
    total = tr_list[-1]
    archetypes = [tr_list[index] for index in spec["archetype_rows"]]
    rows = spec["rows"]
    if len(archetypes) != len(rows):
        raise ValueError("模板样例行数量与表格占位内容数量不一致")
    for tr in tr_list:
        tbl.remove(tr)
    tbl.append(header)
    for archetype, row_texts in zip(archetypes, rows):
        new_row = copy.deepcopy(archetype)
        tbl.append(new_row)
        set_row_texts(table, new_row, row_texts)
    if spec["total"] is not None:
        tbl.append(total)
        total_row = table.rows[-1]
        for cell, text in zip(total_row.cells, spec["total"]):
            set_cell_text(cell, text)


def main() -> int:
    parser = argparse.ArgumentParser(description="生成应急通信装备测试报告固定模板")
    parser.add_argument("--source", required=True, help="原始测试报告 docx 路径")
    parser.add_argument("--output", default=str(DEFAULT_OUTPUT), help="模板输出路径")
    args = parser.parse_args()

    source = pathlib.Path(args.source)
    output = pathlib.Path(args.output)
    if not source.is_file():
        print(f"错误：未找到原始报告文件 {source}", file=sys.stderr)
        return 1
    output.parent.mkdir(parents=True, exist_ok=True)

    document = Document(str(source))
    for index, segments in PARAGRAPH_TEMPLATE.items():
        set_paragraph(document.paragraphs[index], segments)
    for index, spec in TABLE_TEMPLATES.items():
        rebuild_table(document.tables[index], spec)
    document.save(str(output))
    print(f"模板已生成：{output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
